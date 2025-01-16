#!/usr/bin/env python3
from docx import Document
from openpyxl import load_workbook
import os

def update_document_with_worksheet(doc_path, worksheet):
    """Update a single document with values from its corresponding worksheet"""
    if not os.path.exists(doc_path):
        print(f"Warning: Document not found: {doc_path}")
        return
        
    print(f"\nProcessing document: {os.path.basename(doc_path)}")
    
    try:
        # Open the document
        doc = Document(doc_path)
        
        # Get variables from worksheet
        variables = {}
        for row in worksheet.iter_rows(min_row=2):  # Skip header
            if row[0].value:
                var_name = str(row[0].value).strip()
                var_value = str(row[1].value).strip() if row[1].value else ""
                variables[var_name] = var_value
                print(f"Found variable: {var_name} = {var_value}")
        
        changes_made = False
        
        # Update paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            
            for var_name, value in variables.items():
                placeholder = f"[{var_name}]"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, value)
                    print(f"Updated {var_name} in paragraph")
                    changes_made = True
            
            if new_text != original_text:
                para.text = new_text
        
        # Update tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text
                    
                    for var_name, value in variables.items():
                        placeholder = f"[{var_name}]"
                        if placeholder in new_text:
                            new_text = new_text.replace(placeholder, value)
                            print(f"Updated {var_name} in table")
                            changes_made = True
                    
                    if new_text != original_text:
                        cell.text = new_text
        
        if changes_made:
            doc.save(doc_path)
            print(f"✓ Saved updates to {os.path.basename(doc_path)}")
        else:
            print(f"No updates needed for {os.path.basename(doc_path)}")
            
    except Exception as e:
        print(f"Error processing {doc_path}: {str(e)}")

def update_all_documents(excel_path, base_path=None):
    """Process each worksheet and update corresponding documents"""
    if base_path is None:
        base_path = os.getcwd()
    
    print(f"Opening Excel file: {excel_path}")
    wb = load_workbook(excel_path)
    
    # Process each worksheet independently
    for ws in wb.worksheets:
        doc_name = ws.title
        doc_path = os.path.join(base_path, f"{doc_name}.docx")
        update_document_with_worksheet(doc_path, ws)

if __name__ == "__main__":
    update_all_documents("InputFields.xlsx")