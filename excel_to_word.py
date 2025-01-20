#!/usr/bin/env python3
from docx import Document
from openpyxl import load_workbook
import os

def update_text_with_variables(text, variables):
    """Update text content with variables"""
    new_text = text
    made_change = False
    
    for var_name, value in variables.items():
        placeholder = f"[{var_name}]"
        if placeholder in new_text:
            print(f"Found placeholder: {placeholder} -> {value}")
            new_text = new_text.replace(placeholder, value)
            made_change = True
    
    return made_change, new_text

def update_document_with_worksheet(doc_path, worksheet):
    """Update a single document with values from its corresponding worksheet"""
    if not os.path.exists(doc_path):
        print(f"Warning: Document not found: {doc_path}")
        return
        
    print(f"\nProcessing document: {os.path.basename(doc_path)}")
    
    try:
        # Open the document
        doc = Document(doc_path)
        
        # Get variables from worksheet
        variables = {}
        for row in worksheet.iter_rows(min_row=2):  # Skip header
            if row[0].value:
                var_name = str(row[0].value).strip()
                var_value = str(row[1].value).strip() if row[1].value else ""
                variables[var_name] = var_value
                print(f"Found variable: {var_name} = {var_value}")
        
        changes_made = False
        
        # 1. Title Page Processing - Using XML approach
        print("\nProcessing title page elements...")
        xml_body = doc._element.body
        for element in xml_body.iter():
            if element.tag.endswith('}t'):  # Text elements in Word XML
                if element.text:
                    changed, new_text = update_text_with_variables(element.text, variables)
                    if changed:
                        print(f"Updated title page text: '{element.text}' -> '{new_text}'")
                        element.text = new_text
                        changes_made = True
        
        # 2. Regular Paragraph Processing
        print("\nProcessing regular paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            
            # Check if paragraph contains any placeholders
            for var_name, value in variables.items():
                placeholder = f"[{var_name}]"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, value)
                    changes_made = True
            
            # If text changed, update paragraph while preserving formatting
            if new_text != original_text:
                # Clear paragraph content
                p = para._p
                for run in para.runs:
                    p.remove(run._r)
                # Add new text while preserving alignment and style
                new_run = para.add_run(new_text)
                # Copy formatting from first run if it existed
                if len(para.runs) > 0:
                    original_run = para.runs[0]
                    if hasattr(original_run, 'bold'):
                        new_run.bold = original_run.bold
                    if hasattr(original_run, 'italic'):
                        new_run.italic = original_run.italic
                    if hasattr(original_run, 'font'):
                        if hasattr(original_run.font, 'name'):
                            new_run.font.name = original_run.font.name
                        if hasattr(original_run.font, 'size'):
                            new_run.font.size = original_run.font.size
        
        # 3. Table Processing
        print("\nProcessing tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        new_text = original_text
                        
                        # Check if paragraph contains any placeholders
                        for var_name, value in variables.items():
                            placeholder = f"[{var_name}]"
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, value)
                                changes_made = True
                        
                        # If text changed, update paragraph while preserving formatting
                        if new_text != original_text:
                            # Clear paragraph content
                            p = paragraph._p
                            for run in paragraph.runs:
                                p.remove(run._r)
                            # Add new text directly to existing paragraph
                            run = paragraph.add_run(new_text)

        if changes_made:
            doc.save(doc_path)
            print(f"✓ Saved updates to {os.path.basename(doc_path)}")
        else:
            print(f"No updates needed for {os.path.basename(doc_path)}")
            
    except Exception as e:
        print(f"Error processing {doc_path}: {str(e)}")

def update_all_documents(excel_path, base_path=None):
    """Process each worksheet and update corresponding documents"""
    if base_path is None:
        base_path = os.getcwd()
    
    print(f"Opening Excel file: {excel_path}")
    wb = load_workbook(excel_path)
    
    # Process each worksheet independently
    for ws in wb.worksheets:
        doc_name = ws.title
        doc_path = os.path.join(base_path, f"{doc_name}.docx")
        update_document_with_worksheet(doc_path, ws)

if __name__ == "__main__":
    update_all_documents("InputFields.xlsx")